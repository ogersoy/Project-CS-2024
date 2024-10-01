from docx import Document

# Load and read the file
def load_docx(file):
    doc = Document(file)
    content = []  # To store paragraphs and tables
    start_extracting = False  # Flag to start extracting content after 'Scope'
    stop_extracting = False  # Flag to stop extracting when 'Change History' is encountered

    # Loop through the document paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()

        # Check if 'Scope' is found, and start extracting content from this point onwards
        if "Scope" in text:
            start_extracting = True  # Start extracting content from this point
            content.append(text)  # Append 'Scope' to content
            continue

        # If we find 'Change History', stop extracting content and don't append it
        if "Change history" in text:
            stop_extracting = True  # Stop further content extraction
            break  # Stop extracting paragraphs immediately

        # Skip content before 'Scope' or after 'Change History'
        if not start_extracting or stop_extracting:
            continue

        # Append paragraphs after 'Scope'
        if start_extracting and text:
            content.append(text)

    for table in doc.tables:
        table_data=[]
        for row in table.rows:
            if any("Change history" in cell.text for cell in row.cells):
                break  # Stop extracting tables once 'Change History' is detected
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append("| " + " | ".join(row_data) + " |")

        if table_data:
            header = table_data[0]
            separator = "| " + " | ".join(["---"] * len(table_data[0].split("|"))) + " |"
            content.append(header)
            content.append(separator)
            content.extend(table_data[1:])

    return '\n'.join(content)

# print(load_docx(doc))

# Save the output to a txt file
output_txt_path = "output3.txt"
with open(output_txt_path, 'w', encoding='utf-8') as txt_file:
    txt_file.write(load_docx('TSpec-LLM/3GPP-clean/Rel-8/21_series/21101-840.docx'))